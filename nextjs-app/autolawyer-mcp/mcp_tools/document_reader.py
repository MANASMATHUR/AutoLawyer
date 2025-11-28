from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List

from docx import Document
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    name: str
    path: str
    content: str
    metadata: Dict


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _load_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_text,
}


def ingest_documents(files: Iterable[Dict]) -> List[Dict]:
    """
    Accepts list of dicts {name, path or bytes}. Returns normalized metadata.
    """
    outputs: List[Dict] = []
    for raw in files:
        path = Path(raw["path"]).expanduser()
        if not path.exists():
            raise FileNotFoundError(path)
        ext = path.suffix.lower()
        loader = LOADERS.get(ext)
        if not loader:
            raise ValueError(f"Unsupported extension {ext}")

        content = loader(path)
        metadata = {
            "size": path.stat().st_size,
            "extension": ext,
        }
        outputs.append(
            LoadedDocument(
                name=raw.get("name", path.name),
                path=str(path),
                content=content,
                metadata=metadata,
            ).__dict__
        )
    return outputs


